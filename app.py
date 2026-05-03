"""
app.py — Streamlit Frontend für den KI-Co-Portfolio-Manager
Berner Fachhochschule | Bachelor Thesis 2025/26 | Luca Lüdi

Start: streamlit run app.py
"""

import streamlit as st
import pandas as pd
import json
import os
import sys
import time
from datetime import datetime
from pathlib import Path

from tools.finance_tools import search_ticker

# ── Page Config (muss als erstes Streamlit-Call stehen) ──────────────────────
st.set_page_config(
    page_title="KI-Co-Portfolio-Manager",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Custom CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
  @import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;600;700&family=DM+Sans:wght@300;400;500&display=swap');

  html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
  }

  /* Header */
  .main-header {
    background: linear-gradient(135deg, #0f1923 0%, #1a2f45 100%);
    padding: 2rem 2.5rem;
    border-radius: 12px;
    margin-bottom: 2rem;
    border-left: 4px solid #c9a84c;
  }
  .main-header h1 {
    font-family: 'Playfair Display', serif;
    color: #f0e6d0;
    font-size: 2rem;
    margin: 0;
    letter-spacing: -0.5px;
  }
  .main-header p {
    color: #8a9bb0;
    margin: 0.3rem 0 0 0;
    font-size: 0.9rem;
    font-weight: 300;
  }

  /* Empfehlung Badge */
  .badge-kaufen {
    background: #0d4f2e;
    color: #4ade80;
    border: 1px solid #4ade80;
    padding: 0.4rem 1.2rem;
    border-radius: 20px;
    font-weight: 600;
    font-size: 1rem;
    display: inline-block;
    letter-spacing: 1px;
  }
  .badge-halten {
    background: #3d3200;
    color: #fbbf24;
    border: 1px solid #fbbf24;
    padding: 0.4rem 1.2rem;
    border-radius: 20px;
    font-weight: 600;
    font-size: 1rem;
    display: inline-block;
    letter-spacing: 1px;
  }
  .badge-verkaufen {
    background: #4f0d0d;
    color: #f87171;
    border: 1px solid #f87171;
    padding: 0.4rem 1.2rem;
    border-radius: 20px;
    font-weight: 600;
    font-size: 1rem;
    display: inline-block;
    letter-spacing: 1px;
  }

  /* Metric Cards */
  .metric-card {
    background: #f8f9fa;
    border: 1px solid #e9ecef;
    border-radius: 10px;
    padding: 1.2rem;
    text-align: center;
  }
  .metric-card .label {
    font-size: 0.75rem;
    color: #6c757d;
    text-transform: uppercase;
    letter-spacing: 0.8px;
    font-weight: 500;
    margin-bottom: 0.3rem;
  }
  .metric-card .value {
    font-size: 1.5rem;
    font-weight: 600;
    color: #1a2f45;
    font-family: 'Playfair Display', serif;
  }
  .metric-card .sub {
    font-size: 0.8rem;
    color: #6c757d;
    margin-top: 0.2rem;
  }

  /* Section Headers */
  .section-header {
    font-family: 'Playfair Display', serif;
    color: #1a2f45;
    font-size: 1.1rem;
    font-weight: 600;
    border-bottom: 2px solid #c9a84c;
    padding-bottom: 0.4rem;
    margin: 1.5rem 0 1rem 0;
  }

  /* Ampel */
  .ampel-positiv { color: #4ade80; font-size: 1.2rem; }
  .ampel-neutral  { color: #fbbf24; font-size: 1.2rem; }
  .ampel-negativ  { color: #f87171; font-size: 1.2rem; }

  /* Quality Check */
  .check-bestanden   { color: #4ade80; }
  .check-warnung     { color: #fbbf24; }
  .check-fehlgeschlagen { color: #f87171; }

  /* Investment Case Bullets */
  .inv-bullet {
    background: #f8f9fa;
    border-left: 3px solid #c9a84c;
    padding: 0.7rem 1rem;
    margin-bottom: 0.5rem;
    border-radius: 0 6px 6px 0;
    font-size: 0.9rem;
    line-height: 1.5;
    color: #212529;
  }

  /* Scenario Cards */
  .scenario-bear {
    background: #fff5f5;
    border: 1px solid #fca5a5;
    border-radius: 8px;
    padding: 1rem;
    color: #212529;
  }
  .scenario-base {
    background: #fffbeb;
    border: 1px solid #fcd34d;
    border-radius: 8px;
    padding: 1rem;
    color: #212529;
  }
  .scenario-bull {
    background: #f0fdf4;
    border: 1px solid #86efac;
    border-radius: 8px;
    padding: 1rem;
    color: #212529;
  }

  /* Conviction Killer */
  .conviction-killer {
    background: #fff5f5;
    border: 1px solid #f87171;
    border-radius: 8px;
    padding: 1rem;
    margin-bottom: 0.8rem;
    color: #212529;
  }

  /* Sidebar */
  .sidebar-info {
    background: #f8f9fa;
    border-radius: 8px;
    padding: 1rem;
    font-size: 0.85rem;
    color: #6c757d;
    margin-top: 1rem;
  }

  /* Disclaimer */
  .disclaimer {
    background: #f8f9fa;
    border: 1px solid #dee2e6;
    border-radius: 8px;
    padding: 1rem;
    font-size: 0.78rem;
    color: #6c757d;
    margin-top: 2rem;
    line-height: 1.6;
  }

  /* Upside positiv/negativ */
  .upside-pos { color: #4ade80; font-weight: 600; }
  .upside-neg { color: #f87171; font-weight: 600; }

  /* Monitoring Checkbox */
  .monitor-item {
    display: flex;
    align-items: flex-start;
    gap: 0.6rem;
    padding: 0.5rem 0;
    border-bottom: 1px solid #f0f0f0;
    font-size: 0.88rem;
  }

  /* Hide Streamlit branding */
  #MainMenu {visibility: hidden;}
  footer {visibility: hidden;}
  .stDeployButton {display: none;}
</style>
""", unsafe_allow_html=True)


# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def get_badge_html(recommendation: str) -> str:
    rec = recommendation.upper()
    styles = {
        "KAUFEN":
            "background:#0d4f2e; color:#4ade80; border:1px solid #4ade80;",
        "ÜBERGEWICHTEN":
            "background:#1a3d1a; color:#86efac; border:1px solid #86efac;",
        "HALTEN":
            "background:#3d3200; color:#fbbf24; border:1px solid #fbbf24;",
        "UNTERGEWICHTEN":
            "background:#3d1f00; color:#fb923c; border:1px solid #fb923c;",
        "VERKAUFEN":
            "background:#4f0d0d; color:#f87171; border:1px solid #f87171;",
    }
    style = styles.get(rec, styles["HALTEN"])
    return (f'<span style="padding:0.4rem 1.2rem; border-radius:20px; '
            f'font-weight:600; font-size:1rem; display:inline-block; '
            f'letter-spacing:1px; {style}">{rec}</span>')


def get_ampel_icon(signal: str) -> str:
    return {
        "positiv":  "🟢",
        "neutral":  "🟡",
        "negativ":  "🔴",
        "tailwind": "🟢",
        "headwind": "🔴",
    }.get(signal.lower(), "🟡")


def get_upside_html(upside: float) -> str:
    if upside > 0:
        return f'<span class="upside-pos">▲ +{upside:.1f}%</span>'
    return f'<span class="upside-neg">▼ {upside:.1f}%</span>'


def format_number(value, prefix="", suffix="", decimals=2) -> str:
    try:
        v = float(value)
        return f"{prefix}{v:,.{decimals}f}{suffix}"
    except (TypeError, ValueError):
        return str(value) if value else "n/v"


def load_demo_output() -> dict | None:
    """Lädt gespeicherte Demo-Outputs falls vorhanden."""
    for fname in ["output_memo_HOLN_SW.json", "output_memo.json"]:
        if Path(fname).exists():
            with open(fname, encoding="utf-8") as f:
                return json.load(f)
    return None


def _build_word_memo(data: dict, ticker: str, date: str, ccy: str) -> bytes:
    """Erstellt ein Word-Dokument (.docx) aus den Investment-Memo-Daten."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor(0x1a, 0x2f, 0x45)

    def add_body(text):
        p = doc.add_paragraph(text)
        p.runs[0].font.size = Pt(10) if p.runs else None

    # Titel
    title = doc.add_heading(f"Investment Memo — {ticker}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Datum: {date} | KI-Co-Portfolio-Manager (BFH 2025/26, Luca Lüdi)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Empfehlung & Kursziel
    rec = data.get("recommendation", "n/v")
    pt = data.get("price_target", 0)
    cp = data.get("current_price", 0)
    add_heading("Empfehlung & Kursziel", 1)
    p = doc.add_paragraph()
    p.add_run(f"Empfehlung: ").bold = True
    p.runs[-1].bold = True
    p.add_run(f"{rec}   |   Kursziel: {ccy} {pt:.2f}   |   Aktueller Kurs: {ccy} {cp:.2f}")

    # Unternehmensbeschreibung
    add_heading("Unternehmensbeschreibung", 1)
    doc.add_paragraph(data.get("company_description", "n/v"))

    # Investment Case
    add_heading("Investment Case", 1)
    for point in data.get("investment_case", []):
        doc.add_paragraph(point, style="List Bullet")

    # Finale Begründung
    add_heading("Finale Begründung", 1)
    doc.add_paragraph(data.get("final_reasoning", "n/v"))

    # Szenarien
    add_heading("Szenarien", 1)
    scenarios = data.get("scenarios", [])
    if scenarios:
        table = doc.add_table(rows=1, cols=len(scenarios))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, s in enumerate(scenarios):
            hdr[i].text = s.get("name", "")
            hdr[i].paragraphs[0].runs[0].bold = True
        row = table.add_row().cells
        for i, s in enumerate(scenarios):
            row[i].text = (
                f"Kursziel: {ccy} {s.get('price_target', 0):.2f}\n"
                f"Wahrsch.: {s.get('probability_pct', 0)}%\n"
                f"Kernannahme: {s.get('key_assumption', '')}\n"
                f"Trigger: {s.get('trigger', '')}"
            )

    # Quantifizierte Risiken
    add_heading("Quantifizierte Risiken", 1)
    for risk in data.get("key_risks", []):
        doc.add_paragraph(f"⚠ {risk}", style="List Bullet")

    # Conviction Killers
    add_heading("Conviction Killers", 1)
    for ck in data.get("conviction_killers", []):
        desc = ck.get("description", "") if isinstance(ck, dict) else str(ck)
        monitor = ck.get("monitoring_indicator", "") if isinstance(ck, dict) else ""
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(desc).bold = True
        if monitor:
            p.add_run(f" → Monitor: {monitor}")

    # Disclaimer
    doc.add_paragraph()
    disc = doc.add_paragraph(
        "Disclaimer: Dieses Dokument wurde automatisch durch das KI-Co-Portfolio-Manager System "
        "generiert und dient ausschliesslich zu Forschungs- und Demonstrationszwecken. "
        "Es stellt keine Anlageberatung dar."
    )
    disc.runs[0].font.size = Pt(8)
    disc.runs[0].font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Pipeline Import (mit Fehlerbehandlung) ────────────────────────────────────

def run_analysis(ticker: str) -> dict:
    """Führt die Pipeline aus und gibt das Ergebnis zurück."""
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from graph.supervisor import run_supervisor
    return run_supervisor(ticker)


# ── Sidebar ───────────────────────────────────────────────────────────────────

with st.sidebar:
    st.markdown("""
    <div style="text-align:center; padding: 1rem 0;">
      <div style="font-family:'Playfair Display',serif; font-size:1.3rem; 
                  color:#1a2f45; font-weight:700;">
        📊 Portfolio Manager
      </div>
      <div style="font-size:0.75rem; color:#8a9bb0; margin-top:0.2rem;">
        KI-Co-Portfolio-Manager
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.divider()

    st.markdown("### Aktie analysieren")

    company_query = st.text_input(
        "Unternehmen suchen",
        placeholder="z.B. Holcim, Apple, Nestlé, Roche...",
        help="Firmenname oder Ticker eingeben"
    )

    if len(company_query) >= 2 and company_query != st.session_state.get("last_query", ""):
        with st.spinner("Suche..."):
            st.session_state.search_results = search_ticker(company_query)
            st.session_state.last_query = company_query

    search_results = st.session_state.get("search_results", [])

    if search_results:
        st.markdown("**Ergebnisse:**")
        options = [r["display"] for r in search_results]
        options_map = {r["display"]: r["ticker"] for r in search_results}
        selected_display = st.radio(
            label="Aktie auswählen",
            options=options,
            label_visibility="collapsed",
        )
        if selected_display:
            st.session_state.selected_ticker = options_map[selected_display]
            st.success(f"✓ {st.session_state.selected_ticker}")
    elif len(company_query) >= 2:
        st.warning("Keine Ergebnisse — versuche es mit dem Ticker direkt")
        manual = st.text_input("Ticker direkt eingeben", placeholder="z.B. HOLN.SW")
        if manual:
            st.session_state.selected_ticker = manual.upper().strip()

    selected_ticker = st.session_state.get("selected_ticker", "")

    st.divider()

    run_button = st.button(
        "▶ Analyse starten",
        type="primary",
        use_container_width=True,
        disabled=not selected_ticker,
    )

    st.divider()
    st.markdown("**Beispiele:**")

    examples = [
        ("Holcim", "HOLN.SW"),
        ("Nestlé", "NESN.SW"),
        ("Apple", "AAPL"),
        ("Microsoft", "MSFT"),
        ("Novartis", "NOVN.SW"),
        ("Roche", "ROG.SW"),
    ]

    cols = st.columns(2)
    for i, (name, ticker_ex) in enumerate(examples):
        with cols[i % 2]:
            if st.button(name, use_container_width=True, key=f"ex_{ticker_ex}"):
                st.session_state.selected_ticker = ticker_ex
                st.session_state.search_results = []
                st.session_state.last_query = ""
                st.rerun()

    # Demo-Modus
    st.divider()
    demo_button = st.button(
        "📂 Demo laden (Holcim)",
        use_container_width=True,
        help="Lädt einen gespeicherten Demo-Output ohne API-Calls"
    )

    st.markdown("""
    <div class="sidebar-info">
      <strong>System-Info</strong><br>
      3 KI-Agenten:<br>
      • Fundamental-Agent<br>
      • News/Sentiment-Agent<br>
      • Risk/Advocatus-Agent<br><br>
      Dauer: ca. 60–90 Sek.<br>
      Modell: Claude Sonnet + GPT-4o
    </div>
    """, unsafe_allow_html=True)


# ── Main Header ───────────────────────────────────────────────────────────────

st.markdown("""
<div class="main-header">
  <h1>KI-Co-Portfolio-Manager</h1>
  <p>Berner Fachhochschule · Bachelor Thesis 2025/26 · Luca Lüdi · 
     Multi-Agent Investment Analysis System</p>
</div>
""", unsafe_allow_html=True)


# ── Session State ─────────────────────────────────────────────────────────────

if "result" not in st.session_state:
    st.session_state.result = None
if "ticker" not in st.session_state:
    st.session_state.ticker = None


# ── Demo laden ────────────────────────────────────────────────────────────────

if demo_button:
    demo = load_demo_output()
    if demo:
        st.session_state.result = demo
        st.session_state.ticker = demo.get("ticker", "HOLN.SW")
        st.success("Demo-Output geladen (Holcim AG, 29.04.2026)")
    else:
        st.warning("Kein gespeicherter Demo-Output gefunden. Bitte zuerst eine Analyse ausführen.")


# ── Analyse starten ───────────────────────────────────────────────────────────

if run_button and selected_ticker:
    ticker = selected_ticker.upper().strip()
    st.session_state.ticker = ticker

    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from agents.fundamental_agent import run_fundamental_agent
    from agents.news_agent import run_news_agent
    from agents.risk_agent import run_risk_agent
    from graph.supervisor import synthesize_memo

    try:
        with st.status(f"Analyse läuft für **{ticker}**…", expanded=True) as status:

            # ── Schritt 1: Fundamental-Agent ─────────────────────────────────
            st.write("🔍 **Fundamental-Agent** — Daten abrufen & DCF-Bewertung…")
            fundamental_output = run_fundamental_agent(ticker)
            rec   = fundamental_output.get("recommendation", "n/v") if isinstance(fundamental_output, dict) else getattr(fundamental_output, "recommendation", "n/v")
            fv    = fundamental_output.get("fair_value_estimate", "n/v") if isinstance(fundamental_output, dict) else getattr(fundamental_output, "fair_value_estimate", "n/v")
            updn  = fundamental_output.get("upside_downside_pct", "n/v") if isinstance(fundamental_output, dict) else getattr(fundamental_output, "upside_downside_pct", "n/v")
            val   = fundamental_output.get("valuation_assessment", "n/v") if isinstance(fundamental_output, dict) else getattr(fundamental_output, "valuation_assessment", "n/v")
            st.write(f"   ✅ Empfehlung: **{rec}** | Fair Value: {fv} | Upside: {updn}% | Bewertung: {val}")

            # Investment Case Bullets
            ic = fundamental_output.get("investment_case", []) if isinstance(fundamental_output, dict) else getattr(fundamental_output, "investment_case", [])
            if ic:
                for pt in ic[:3]:
                    point_txt = pt.get("point", str(pt)) if isinstance(pt, dict) else str(pt)
                    st.write(f"   • {point_txt}")

            # ── Schritt 2: News/Sentiment-Agent ──────────────────────────────
            st.write("📰 **News/Sentiment-Agent** — Makro & Nachrichten auswerten…")
            fundamental_context = (
                f"Empfehlung: {rec}, Fair Value: {fv}, Bewertung: {val}"
            )
            news_output = run_news_agent(ticker, fundamental_context)
            sentiment    = news_output.get("overall_sentiment_score", "n/v") if isinstance(news_output, dict) else getattr(news_output, "overall_sentiment_score", "n/v")
            st_outlook   = news_output.get("short_term_outlook", "n/v") if isinstance(news_output, dict) else getattr(news_output, "short_term_outlook", "n/v")
            macro_dir    = news_output.get("overall_macro_direction", "n/v") if isinstance(news_output, dict) else getattr(news_output, "overall_macro_direction", "n/v")
            st.write(f"   ✅ Sentiment: **{sentiment}/10** | Kurzfrist-Outlook: {st_outlook} | Makro: {macro_dir}")

            # Top-News
            news_items = news_output.get("news_items", []) if isinstance(news_output, dict) else getattr(news_output, "news_items", [])
            for ni in news_items[:2]:
                headline = ni.get("headline", "") if isinstance(ni, dict) else getattr(ni, "headline", "")
                src      = ni.get("source", "") if isinstance(ni, dict) else getattr(ni, "source", "")
                impact   = ni.get("sentiment_impact", "") if isinstance(ni, dict) else getattr(ni, "sentiment_impact", "")
                if headline:
                    st.write(f"   • [{impact}] {headline} — *{src}*")

            # ── Schritt 3: Risk/Advocatus-Agent ──────────────────────────────
            st.write("⚖️ **Risk-Agent** — Advocatus Diaboli & Szenarien…")
            risk_output  = run_risk_agent(ticker, fundamental_output, news_output)
            risk_rec     = risk_output.get("original_recommendation", "n/v") if isinstance(risk_output, dict) else getattr(risk_output, "original_recommendation", "n/v")
            counter      = risk_output.get("counter_position", "") if isinstance(risk_output, dict) else getattr(risk_output, "counter_position", "")
            st.write(f"   ✅ Gegenposition zu **{risk_rec}**: {counter}")

            scenarios = risk_output.get("scenarios", []) if isinstance(risk_output, dict) else getattr(risk_output, "scenarios", [])
            for sc in scenarios:
                sname  = sc.get("name", "") if isinstance(sc, dict) else sc.name
                sprob  = sc.get("probability_pct", "") if isinstance(sc, dict) else sc.probability_pct
                spt    = sc.get("price_target", "") if isinstance(sc, dict) else sc.price_target
                st.write(f"   • {sname}: {sprob}% Wahrscheinlichkeit | Kursziel: {spt}")

            ck_list = risk_output.get("conviction_killers", []) if isinstance(risk_output, dict) else getattr(risk_output, "conviction_killers", [])
            for ck in ck_list:
                desc = ck.get("description", "") if isinstance(ck, dict) else getattr(ck, "description", "")
                st.write(f"   🚨 Conviction Killer: {desc}")

            # ── Schritt 4: Qualitätsprüfung ───────────────────────────────────
            st.write("🔎 **Qualitätsprüfung** — Konsistenz der drei Analysen…")
            from graph.supervisor import _build_quality_checks
            quality_checks = _build_quality_checks(fundamental_output, news_output, risk_output)
            ok  = sum(1 for c in quality_checks if c["result"] == "bestanden")
            wrn = sum(1 for c in quality_checks if c["result"] == "Warnung")
            err = sum(1 for c in quality_checks if c["result"] == "fehlgeschlagen")
            st.write(f"   ✅ {ok} bestanden · ⚠️ {wrn} Warnungen · ❌ {err} fehlgeschlagen")
            for c in quality_checks:
                icon_map = {"bestanden": "✅", "Warnung": "⚠️", "fehlgeschlagen": "❌"}
                st.write(f"   {icon_map.get(c['result'], 'ℹ️')} {c['check']}: {c['comment']}")

            # ── Schritt 5: Supervisor-Synthese ────────────────────────────────
            st.write("✍️ **Supervisor** — Finales Investment Memo wird synthetisiert…")
            result = synthesize_memo(ticker, fundamental_output, news_output, risk_output)
            final_rec = result.get("final_recommendation", "n/v")
            conviction = result.get("conviction_level", "n/v")
            pt         = result.get("price_target", "n/v")
            st.write(f"   ✅ Finale Empfehlung: **{final_rec}** | Conviction: {conviction} | Kursziel: {pt}")

            status.update(label=f"✓ Analyse abgeschlossen — {result.get('company', ticker)}", state="complete", expanded=False)

    except Exception as e:
        st.error(f"Fehler bei der Analyse: {str(e)}")
        st.info("Tipp: Prüfen Sie ob der Ticker korrekt ist (z.B. HOLN.SW für Holcim)")


# ── Ergebnis anzeigen ─────────────────────────────────────────────────────────

if st.session_state.result:
    data = st.session_state.result
    rec   = data.get("final_recommendation", "HALTEN")
    conv  = data.get("conviction_level", "n/v")
    pt    = data.get("price_target", 0)
    price = data.get("current_price", 0)
    updn  = data.get("upside_downside_pct", 0)
    ccy   = data.get("currency", "")
    company = data.get("company", "")
    ticker  = data.get("ticker", "")
    date    = data.get("date", datetime.now().strftime("%Y-%m-%d"))
    score   = data.get("data_consistency_score", 0)

    # ── Hero Section ──────────────────────────────────────────────────────────
    st.markdown(f"""
    <div style="background:white; border:1px solid #e9ecef; border-radius:12px; 
                padding:1.5rem 2rem; margin-bottom:1.5rem;
                box-shadow: 0 2px 8px rgba(0,0,0,0.06);">
      <div style="display:flex; justify-content:space-between; align-items:flex-start; 
                  flex-wrap:wrap; gap:1rem;">
        <div>
          <div style="font-family:'Playfair Display',serif; font-size:1.8rem; 
                      color:#1a2f45; font-weight:700;">{company}</div>
          <div style="color:#8a9bb0; font-size:0.9rem; margin-top:0.2rem;">
            {ticker} · {data.get('sector', '')} · {date}
          </div>
          <div style="margin-top:0.8rem;">
            {get_badge_html(rec)}
            <span style="margin-left:0.8rem; color:#6c757d; font-size:0.9rem;">
              Conviction: <strong>{conv}</strong>
            </span>
          </div>
        </div>
        <div style="text-align:right;">
          <div style="font-size:0.8rem; color:#8a9bb0; text-transform:uppercase; 
                      letter-spacing:0.8px;">Konsistenz-Score</div>
          <div style="font-size:2.5rem; font-family:'Playfair Display',serif; 
                      color:#1a2f45; font-weight:700; line-height:1;">
            {score}<span style="font-size:1rem; color:#8a9bb0;">/10</span>
          </div>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Key Metrics ───────────────────────────────────────────────────────────
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.markdown(f"""
        <div class="metric-card">
          <div class="label">Aktueller Kurs</div>
          <div class="value">{ccy} {format_number(price)}</div>
        </div>""", unsafe_allow_html=True)

    with col2:
        st.markdown(f"""
        <div class="metric-card">
          <div class="label">Price Target</div>
          <div class="value">{ccy} {format_number(pt)}</div>
        </div>""", unsafe_allow_html=True)

    with col3:
        upside_html = get_upside_html(float(updn) if updn else 0)
        st.markdown(f"""
        <div class="metric-card">
          <div class="label">Upside / Downside</div>
          <div class="value">{upside_html}</div>
        </div>""", unsafe_allow_html=True)

    with col4:
        conviction_stars = {"hoch": "★★★", "mittel": "★★☆", "niedrig": "★☆☆"}.get(conv, "★☆☆")
        st.markdown(f"""
        <div class="metric-card">
          <div class="label">Conviction</div>
          <div class="value" style="font-size:1.3rem;">{conviction_stars}</div>
          <div class="sub">{conv}</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Tabs ──────────────────────────────────────────────────────────────────
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📋 Investment Memo",
        "📊 Bewertung",
        "📰 Makro & Sentiment",
        "⚠️ Risiken & Szenarien",
        "✅ Qualitätsprüfung"
    ])

    # ────────────────────────────────────────────────────────────────────────
    # TAB 1: Investment Memo
    # ────────────────────────────────────────────────────────────────────────
    with tab1:
        col_left, col_right = st.columns([3, 2])

        with col_left:
            # Unternehmensbeschreibung
            st.markdown('<div class="section-header">Unternehmensbeschreibung</div>',
                       unsafe_allow_html=True)
            st.markdown(data.get("company_description", "n/v"))

            # Investment Case
            st.markdown('<div class="section-header">Investment Case</div>',
                       unsafe_allow_html=True)
            for point in data.get("investment_case", []):
                st.markdown(f'<div class="inv-bullet">{point}</div>',
                           unsafe_allow_html=True)

            # Finale Begründung
            st.markdown('<div class="section-header">Finale Begründung</div>',
                       unsafe_allow_html=True)
            st.markdown(f"""
            <div style="background:#f8f9fa; border-radius:8px; padding:1rem; 
                        font-size:0.88rem; line-height:1.7; color:#495057;">
              {data.get('final_reasoning', 'n/v')}
            </div>""", unsafe_allow_html=True)

        with col_right:
            # Makro-Ampel
            st.markdown('<div class="section-header">Makro-Ampel</div>',
                       unsafe_allow_html=True)
            for amp in data.get("macro_ampel", []):
                icon = get_ampel_icon(amp.get("signal", "neutral"))
                st.markdown(f"""
                <div style="display:flex; gap:0.8rem; padding:0.6rem 0; 
                            border-bottom:1px solid #f0f0f0;">
                  <span style="font-size:1.2rem; flex-shrink:0;">{icon}</span>
                  <div>
                    <div style="font-weight:600; font-size:0.85rem; 
                                color:#1a2f45;">{amp.get('category', '')}</div>
                    <div style="font-size:0.82rem; color:#6c757d; margin-top:0.2rem;">
                      {amp.get('key_point', '')}
                    </div>
                  </div>
                </div>""", unsafe_allow_html=True)

            # Advocatus Diaboli
            st.markdown('<div class="section-header">Advocatus Diaboli</div>',
                       unsafe_allow_html=True)
            st.markdown(f"""
            <div style="background:#fff5f5; border:1px solid #fca5a5; 
                        border-radius:8px; padding:1rem; font-size:0.85rem; 
                        line-height:1.6; color:#7f1d1d;">
              {data.get('advocatus_diaboli_summary', 'n/v')}
            </div>""", unsafe_allow_html=True)

            # Quellen
            st.markdown('<div class="section-header">Quellen</div>',
                       unsafe_allow_html=True)
            for src in data.get("sources", []):
                st.markdown(f"<div style='font-size:0.8rem; color:#6c757d; "
                           f"padding:0.2rem 0;'>• {src}</div>",
                           unsafe_allow_html=True)

    # ────────────────────────────────────────────────────────────────────────
    # TAB 2: Bewertung
    # ────────────────────────────────────────────────────────────────────────
    with tab2:
        # Bewertungstabelle
        st.markdown('<div class="section-header">Bewertungs-Multiples</div>',
                   unsafe_allow_html=True)

        vt = data.get("valuation_table", [])
        if vt:
            rows = []
            for row in vt:
                assessment = row.get("assessment", "FAIR")
                assess_label = {
                    "ELEVATED": "🔴 ELEVATED",
                    "FAIR":     "🟢 FAIR",
                    "DISCOUNT": "🔵 DISCOUNT",
                }.get(assessment, assessment)
                rows.append({
                    "Kennzahl":     row.get("metric", ""),
                    "Aktuell":      row.get("current_value", "n/v"),
                    "Peer Ø":       row.get("peer_average", "n/v"),
                    "Hist. Ø":      row.get("historical_average", "n/v"),
                    "Einschätzung": assess_label,
                    "Quelle":       row.get("source", ""),
                })
            df = pd.DataFrame(rows)
            st.dataframe(
                df,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Kennzahl":     st.column_config.TextColumn(width="medium"),
                    "Aktuell":      st.column_config.TextColumn(width="small"),
                    "Peer Ø":       st.column_config.TextColumn(width="small"),
                    "Hist. Ø":      st.column_config.TextColumn(width="small"),
                    "Einschätzung": st.column_config.TextColumn(width="small"),
                    "Quelle":       st.column_config.TextColumn(width="medium"),
                }
            )
            st.caption("🔴 ELEVATED = über Peer/Hist. Durchschnitt  |  "
                       "🟢 FAIR = im normalen Bereich  |  "
                       "🔵 DISCOUNT = unter Durchschnitt")
        else:
            st.info("Keine Bewertungsdaten verfügbar.")

        # Konsensschätzungen
        st.markdown('<div class="section-header">Konsensschätzungen</div>',
                   unsafe_allow_html=True)
        ce = data.get("consensus_estimates", [])
        if ce:
            rows = []
            for row in ce:
                is_estimate = row.get("type") == "E"
                year_label = f"📊 {row.get('year', '')}" if is_estimate else row.get("year", "")
                rows.append({
                    "Jahr":          year_label,
                    "Umsatz (Mrd.)": row.get("revenue_bn", "n/v"),
                    "EBITDA-%":      row.get("ebitda_margin_pct", "n/v"),
                    "EPS":           row.get("eps", "n/v"),
                    "EV/EBITDA":     row.get("ev_ebitda", "n/v"),
                    "KGV":           row.get("pe_ratio", "n/v"),
                    "# Analysten":   row.get("number_of_analysts", "n/v"),
                })
            df_ce = pd.DataFrame(rows)
            st.dataframe(
                df_ce,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Jahr":          st.column_config.TextColumn(width="small"),
                    "Umsatz (Mrd.)": st.column_config.TextColumn(width="small"),
                    "EBITDA-%":      st.column_config.TextColumn(width="small"),
                    "EPS":           st.column_config.TextColumn(width="small"),
                    "EV/EBITDA":     st.column_config.TextColumn(width="small"),
                    "KGV":           st.column_config.TextColumn(width="small"),
                    "# Analysten":   st.column_config.TextColumn(width="small"),
                }
            )
            st.caption("📊 = Schätzwerte (E = Estimate)")

    # ────────────────────────────────────────────────────────────────────────
    # TAB 3: Makro & Sentiment
    # ────────────────────────────────────────────────────────────────────────
    with tab3:
        col_l, col_r = st.columns(2)

        with col_l:
            st.markdown('<div class="section-header">Makro-Ampel Detail</div>',
                       unsafe_allow_html=True)
            for amp in data.get("macro_ampel", []):
                icon = get_ampel_icon(amp.get("signal", "neutral"))
                signal = amp.get("signal", "neutral")
                color = {"positiv": "#f0fdf4", "negativ": "#fff5f5"}.get(signal, "#fffbeb")
                border = {"positiv": "#86efac", "negativ": "#fca5a5"}.get(signal, "#fcd34d")
                st.markdown(f"""
                <div style="background:{color}; border:1px solid {border}; 
                            border-radius:8px; padding:0.8rem 1rem; margin-bottom:0.6rem;">
                  <div style="font-weight:600; font-size:0.9rem;">
                    {icon} {amp.get('category', '')}
                    <span style="float:right; font-size:0.75rem; color:#6c757d; 
                                 font-weight:400;">{signal.upper()}</span>
                  </div>
                  <div style="font-size:0.83rem; color:#495057; margin-top:0.3rem; 
                              line-height:1.5;">
                    {amp.get('key_point', '')}
                  </div>
                </div>""", unsafe_allow_html=True)

        with col_r:
            st.markdown('<div class="section-header">Monitoring Checklist</div>',
                       unsafe_allow_html=True)
            for item in data.get("monitoring_checklist", []):
                st.markdown(f"""
                <div class="monitor-item">
                  <span style="color:#c9a84c; flex-shrink:0;">□</span>
                  <span style="font-size:0.85rem; color:#495057;">{item}</span>
                </div>""", unsafe_allow_html=True)

    # ────────────────────────────────────────────────────────────────────────
    # TAB 4: Risiken & Szenarien
    # ────────────────────────────────────────────────────────────────────────
    with tab4:
        # Szenarien
        st.markdown('<div class="section-header">Szenarien</div>',
                   unsafe_allow_html=True)
        scenarios = data.get("scenarios", [])
        if scenarios:
            cols = st.columns(len(scenarios))
            css_map = {
                "Bear Case": "scenario-bear",
                "Base Case": "scenario-base",
                "Bull Case": "scenario-bull",
            }
            icon_map = {
                "Bear Case": "🐻",
                "Base Case": "⚖️",
                "Bull Case": "🐂",
            }
            for i, (col, scenario) in enumerate(zip(cols, scenarios)):
                name = scenario.get("name", "")
                css = css_map.get(name, "scenario-base")
                icon = icon_map.get(name, "")
                with col:
                    st.markdown(f"""
                    <div class="{css}">
                      <div style="font-weight:700; font-size:0.95rem; margin-bottom:0.5rem;">
                        {icon} {name}
                      </div>
                      <div style="font-size:1.4rem; font-family:'Playfair Display',serif; 
                                  font-weight:700; margin-bottom:0.3rem;">
                        {ccy} {format_number(scenario.get('price_target', 0))}
                      </div>
                      <div style="font-size:0.8rem; color:#6c757d; margin-bottom:0.8rem;">
                        Wahrscheinlichkeit: <strong>{scenario.get('probability_pct', 0)}%</strong>
                      </div>
                      <div style="font-size:0.82rem; line-height:1.5; margin-bottom:0.5rem;">
                        <strong>Kernannahme:</strong><br>
                        {scenario.get('key_assumption', '')}
                      </div>
                      <div style="font-size:0.78rem; color:#6c757d; line-height:1.4;">
                        <strong>Trigger:</strong> {scenario.get('trigger', '')}
                      </div>
                    </div>""", unsafe_allow_html=True)

        # Key Risks
        st.markdown('<div class="section-header">Quantifizierte Risiken</div>',
                   unsafe_allow_html=True)
        for risk in data.get("key_risks", []):
            st.markdown(f"""
            <div style="background:#fff5f5; border-left:3px solid #f87171;
                        padding:0.7rem 1rem; margin-bottom:0.5rem; border-radius:0 6px 6px 0;
                        font-size:0.88rem; line-height:1.5; color:#212529;">
              ⚠️ {risk}
            </div>""", unsafe_allow_html=True)

        # Conviction Killers
        st.markdown('<div class="section-header">Conviction Killers</div>',
                   unsafe_allow_html=True)
        for ck in data.get("conviction_killers", []):
            desc = ck.get("description", "") if isinstance(ck, dict) else str(ck)
            monitor = ck.get("monitoring_indicator", "") if isinstance(ck, dict) else ""
            st.markdown(f"""
            <div class="conviction-killer">
              <div style="font-weight:600; font-size:0.9rem; color:#7f1d1d; 
                          margin-bottom:0.4rem;">
                🚨 {desc}
              </div>
              {f'<div style="font-size:0.82rem; color:#991b1b;">→ Monitor: {monitor}</div>' if monitor else ''}
            </div>""", unsafe_allow_html=True)

    # ────────────────────────────────────────────────────────────────────────
    # TAB 5: Qualitätsprüfung
    # ────────────────────────────────────────────────────────────────────────
    with tab5:
        col_l, col_r = st.columns([2, 1])

        with col_l:
            st.markdown('<div class="section-header">Qualitätschecks</div>',
                       unsafe_allow_html=True)
            checks = data.get("quality_checks", [])
            for check in checks:
                result_val = check.get("result", "")
                icon = {"bestanden": "✅", "Warnung": "⚠️", "fehlgeschlagen": "❌"}.get(result_val, "ℹ️")
                bg = {
                    "bestanden": "#f0fdf4",
                    "Warnung": "#fffbeb",
                    "fehlgeschlagen": "#fff5f5"
                }.get(result_val, "#f8f9fa")
                st.markdown(f"""
                <div style="background:{bg}; border-radius:6px; padding:0.7rem 1rem; 
                            margin-bottom:0.4rem; font-size:0.85rem;">
                  <span style="margin-right:0.5rem;">{icon}</span>
                  <strong>{check.get('check', '')}</strong>
                  <div style="color:#6c757d; margin-top:0.2rem; font-size:0.8rem;">
                    {check.get('comment', '')}
                  </div>
                </div>""", unsafe_allow_html=True)

        with col_r:
            st.markdown('<div class="section-header">Konsistenz-Score</div>',
                       unsafe_allow_html=True)
            score_color = "#4ade80" if score >= 7 else ("#fbbf24" if score >= 5 else "#f87171")
            st.markdown(f"""
            <div style="text-align:center; padding:2rem;">
              <div style="font-size:4rem; font-family:'Playfair Display',serif; 
                          font-weight:700; color:{score_color}; line-height:1;">
                {score}
              </div>
              <div style="color:#6c757d; font-size:0.9rem;">von 10</div>
              <div style="margin-top:1rem; font-size:0.85rem; color:#495057; 
                          line-height:1.6;">
                {data.get('consistency_notes', '')}
              </div>
            </div>""", unsafe_allow_html=True)

    # ── Download Buttons ──────────────────────────────────────────────────────
    st.divider()
    st.markdown("### Export")
    dl_col1, dl_col2, dl_col3, _ = st.columns([1, 1, 1, 2])

    with dl_col1:
        json_str = json.dumps(data, indent=2, ensure_ascii=False, default=str)
        st.download_button(
            label="⬇️ JSON herunterladen",
            data=json_str,
            file_name=f"investment_memo_{ticker}_{date}.json",
            mime="application/json",
            use_container_width=True,
        )

    with dl_col2:
        # Einfaches Text-Memo für Download
        from graph.supervisor import format_investment_memo
        try:
            txt = format_investment_memo(data)
        except Exception:
            txt = json_str
        st.download_button(
            label="⬇️ Text-Memo herunterladen",
            data=txt,
            file_name=f"investment_memo_{ticker}_{date}.txt",
            mime="text/plain",
            use_container_width=True,
        )

    with dl_col3:
        try:
            docx_bytes = _build_word_memo(data, ticker, date, ccy)
            st.download_button(
                label="⬇️ Word herunterladen",
                data=docx_bytes,
                file_name=f"investment_memo_{ticker}_{date}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except ImportError:
            st.warning("python-docx nicht installiert — `pip install python-docx`")

    # ── Disclaimer ────────────────────────────────────────────────────────────
    st.markdown(f"""
    <div class="disclaimer">
      <strong>Disclaimer:</strong> Dieses Dokument wurde automatisch durch das 
      KI-Co-Portfolio-Manager System generiert (Bachelor Thesis BFH 2025/26, Luca Lüdi) 
      und dient ausschliesslich zu Forschungs- und Demonstrationszwecken. 
      Es stellt keine Anlageberatung dar. Alle Angaben basieren auf öffentlich 
      verfügbaren Daten (yfinance, Finnhub, IR-Dokumente) zum Zeitpunkt der Analyse. 
      Eine Haftung für die Richtigkeit der Angaben wird nicht übernommen.
    </div>
    """, unsafe_allow_html=True)


# ── Leerer Zustand ────────────────────────────────────────────────────────────
else:
    st.markdown("""
    <div style="text-align:center; padding:4rem 2rem; color:#8a9bb0;">
      <div style="font-size:3rem; margin-bottom:1rem;">📊</div>
      <div style="font-family:'Playfair Display',serif; font-size:1.5rem; 
                  color:#1a2f45; margin-bottom:0.5rem;">
        Bereit zur Analyse
      </div>
      <div style="font-size:0.95rem; max-width:400px; margin:0 auto; line-height:1.7;">
        Geben Sie links einen Ticker ein und starten Sie die Analyse —
        oder laden Sie den Holcim Demo-Output um das System sofort zu erkunden.
      </div>
      <div style="margin-top:2rem; display:flex; justify-content:center; 
                  gap:2rem; font-size:0.85rem;">
        <div>🤖 3 KI-Agenten</div>
        <div>📄 IR-Dokumente</div>
        <div>📰 Echtzeit-News</div>
        <div>⚖️ Advocatus Diaboli</div>
      </div>
    </div>
    """, unsafe_allow_html=True)